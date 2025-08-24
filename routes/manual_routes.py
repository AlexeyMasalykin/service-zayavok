from flask import Blueprint, render_template, request, jsonify, redirect, url_for, send_file, send_from_directory
from db.models import Session, Application, ApplicationItem, get_moscow_time
import json
from datetime import datetime
import pandas as pd
import io
from docx import Document
from docx.shared import Inches
import tempfile
import os
from ai.analyzer import analyze_with_ai
from utils.smtp_sender import notify_user_direct
from flask import current_app

manual_bp = Blueprint('manual', __name__)

def create_text_for_analysis(application, items):
    """Создает текстовое представление ручной заявки для ИИ-анализа"""
    text = f"""Заявка на закупку

Инициатор заявки: {application.name or 'Не указан'}
Должность: {application.position or 'Не указана'}
Описание объекта закупки: {application.description or 'Не указано'}
Ориентировочная стоимость: {application.cost or 'Не указана'}
Корпус/филиал: {application.department or 'Не указан'}
Общая сумма: {application.total_sum or 'Не указана'}

Обоснование закупки: {application.justification or 'Не указано'}

Таблица товаров:
№ | Наименование товара | Характеристики | Ед. изм. | Кол-во | Цена за ед., руб
"""
    
    if items:
        for i, item in enumerate(items, 1):
            text += f"{i} | {item.item_name or ''} | {item.characteristics or ''} | {item.unit or ''} | {item.quantity or ''} | {item.price or ''}\n"
    else:
        text += "Товары не указаны\n"
    
    return text

@manual_bp.route('/manual')
def manual_form():
    """Страница для ручного заполнения заявки"""
    return render_template('manual_form.html')

@manual_bp.route('/submit_manual', methods=['POST'])
def submit_manual():
    """Обработка ручной заявки"""
    try:
        session = Session()
        
        # Основные поля
        application = Application(
            name=request.form.get('initiator', ''),
            position=request.form.get('position', ''),
            description=request.form.get('description', ''),
            cost=request.form.get('cost', ''),
            phone=request.form.get('phone', ''),
            email=request.form.get('email', ''),
            department=request.form.get('department', ''),
            total_sum=request.form.get('total_sum', ''),
            justification=request.form.get('justification', ''),
            is_manual=True,
            status='новая'
        )
        
        session.add(application)
        session.flush()  # Получаем ID
        
        # Обработка таблицы элементов
        items_list = []
        items_data = request.form.get('items_data', '[]')
        try:
            items = json.loads(items_data)
            for item in items:
                if item.get('item_name'):  # Только если есть название
                    app_item = ApplicationItem(
                        application_id=application.id,
                        item_name=item.get('item_name', ''),
                        characteristics=item.get('characteristics', ''),
                        unit=item.get('unit', ''),
                        quantity=int(item.get('quantity', 0)) if item.get('quantity') else 0,
                        price=float(item.get('price', 0)) if item.get('price') else 0.0
                    )
                    session.add(app_item)
                    items_list.append(app_item)
        except (json.JSONDecodeError, ValueError) as e:
            print(f"Ошибка обработки элементов: {e}")
        
        # Выполняем ИИ-анализ ручной заявки
        try:
            analysis_text = create_text_for_analysis(application, items_list)
            ai_feedback, rating = analyze_with_ai(analysis_text)
            application.ai_analysis = ai_feedback  # Сохраняем ИИ-анализ в отдельное поле
            application.rating = rating     # Сохраняем рейтинг
            print(f"ИИ-анализ выполнен. Рейтинг: {rating}/10")
        except Exception as e:
            print(f"Ошибка ИИ-анализа: {e}")
            application.ai_analysis = "⚠ Не удалось выполнить ИИ-анализ"
            application.rating = 0.0
        
        session.commit()
        
        # Отправляем уведомление пользователю
        if application.email:
            notify_user_direct(current_app, application.email, application.id, application.description)
        
        return jsonify({
            'success': True, 
            'message': f'Заявка #{application.id} успешно создана и проанализирована ИИ!',
            'application_id': application.id,
            'rating': application.rating,
            'redirect_url': f'/manual/result/{application.id}'
        })
        
    except Exception as e:
        session.rollback()
        print(f"Ошибка создания заявки: {e}")
        return jsonify({
            'success': False, 
            'message': f'Ошибка создания заявки: {str(e)}'
        }), 500
    finally:
        session.close()

@manual_bp.route('/manual/result/<int:app_id>')
def manual_result(app_id):
    """Страница результата ручной заявки с ИИ-анализом"""
    try:
        session = Session()
        application = session.query(Application).filter_by(id=app_id, is_manual=True).first()
        
        if not application:
            return "Заявка не найдена", 404
        
        return render_template('manual_result.html', 
                             application=application, 
                             application_id=app_id,
                             datetime=datetime)
        
    except Exception as e:
        print(f"Ошибка отображения результата: {e}")
        return "Ошибка загрузки результата", 500
    finally:
        session.close()

@manual_bp.route('/import_excel', methods=['POST'])
def import_excel():
    """Импорт таблицы из Excel"""
    try:
        if 'excel_file' not in request.files:
            return jsonify({'success': False, 'message': 'Файл не выбран'}), 400
            
        file = request.files['excel_file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Файл не выбран'}), 400
            
        if not file.filename.lower().endswith(('.xlsx', '.xls')):
            return jsonify({'success': False, 'message': 'Поддерживаются только Excel файлы'}), 400
        
        # Читаем Excel файл
        df = pd.read_excel(file, engine='openpyxl' if file.filename.endswith('.xlsx') else 'xlrd')
        
        # Преобразуем в список для фронтенда
        items = []
        for _, row in df.iterrows():
            try:
                # Безопасное преобразование данных
                item_name = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ''
                characteristics = str(row.iloc[1]).strip() if len(row) > 1 and pd.notna(row.iloc[1]) else ''
                unit = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ''
                
                # Обработка количества
                try:
                    quantity = int(float(row.iloc[3])) if len(row) > 3 and pd.notna(row.iloc[3]) else 0
                except (ValueError, TypeError):
                    quantity = 0
                
                # Обработка цены
                try:
                    price = float(row.iloc[4]) if len(row) > 4 and pd.notna(row.iloc[4]) else 0.0
                except (ValueError, TypeError):
                    price = 0.0
                
                item = {
                    'item_name': item_name,
                    'characteristics': characteristics,
                    'unit': unit,
                    'quantity': quantity,
                    'price': price
                }
                
                # Добавляем только если есть название товара
                if item_name and item_name.lower() not in ['nan', 'none', '']:
                    items.append(item)
                    
            except Exception as e:
                print(f"Ошибка обработки строки Excel: {e}")
                continue
        
        return jsonify({'success': True, 'items': items})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Ошибка импорта: {str(e)}'}), 500

@manual_bp.route('/export_word/<int:app_id>')
def export_word(app_id):
    """Экспорт заявки в Word по форме"""
    try:
        session = Session()
        application = session.query(Application).filter_by(id=app_id).first()
        
        if not application:
            return "Заявка не найдена", 404
        
        # Создаем Word документ
        doc = Document()
        
        # Заголовок
        header = doc.add_paragraph()
        header.alignment = 2  # Выравнивание по правому краю
        header.add_run("Директору ГБПОУ НО НМК\n\nВ.Н. Гречко")
        
        doc.add_paragraph()  # Пустая строка
        
        # Основные поля
        doc.add_paragraph(f"Инициатор заявки (ФИО, должность): {application.name or ''}")
        doc.add_paragraph(f"Описание объекта закупки (кратко): {application.description or ''}")
        doc.add_paragraph(f"Ориентировочная стоимость: {application.cost or ''}")
        doc.add_paragraph(f"Номер телефона инициатора заявки: {application.phone or ''}")
        doc.add_paragraph(f"Адрес электронной почты: {application.email or ''}")
        doc.add_paragraph(f"Корпус, филиал: {application.department or ''}")
        
        doc.add_paragraph()
        doc.add_paragraph("Объект закупки:")
        
        # Таблица
        if application.items:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Заголовки
            headers = ['№ п/п', 'Наименование товара', 'Характеристики', 'Ед. изм.', 'Кол-во', 'Цена за ед., руб']
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
            
            # Данные
            for i, item in enumerate(application.items, 1):
                row = table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = item.item_name or ''
                row.cells[2].text = item.characteristics or ''
                row.cells[3].text = item.unit or ''
                row.cells[4].text = str(item.quantity) if item.quantity else ''
                row.cells[5].text = str(item.price) if item.price else ''
        
        doc.add_paragraph()
        doc.add_paragraph(f"Общая сумма: {application.total_sum or ''}")
        
        doc.add_paragraph()
        doc.add_paragraph("Обоснование закупки:")
        doc.add_paragraph(application.justification or '')
        
        # Сохраняем во временный файл
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        filename = f"zayavka_{app_id}_{get_moscow_time().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Ошибка экспорта: {e}")
        return f"Ошибка экспорта: {str(e)}", 500
    finally:
        session.close()
        # Удаляем временный файл
        try:
            if 'temp_file' in locals():
                os.unlink(temp_file.name)
        except OSError:
            pass

@manual_bp.route('/download_excel_template')
def download_excel_template():
    """Скачивание образца Excel файла"""
    try:
        # Путь к образцу файла
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'test_table.xlsx')
        
        if not os.path.exists(template_path):
            return "Образец файла не найден", 404
            
        return send_file(
            template_path,
            as_attachment=True,
            download_name='образец_таблицы_товаров.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        return f"Ошибка скачивания образца: {str(e)}", 500
