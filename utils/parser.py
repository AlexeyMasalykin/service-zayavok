from docx import Document
import io
# import fitz  # PDF support removed
import re

def extract_text(file_bytes: bytes, ext: str) -> str:
    text = ""

    if ext == 'docx':
        doc = Document(io.BytesIO(file_bytes))

        # 📄 Читаем обычные абзацы
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text += "\n".join(paragraphs)

        # 📊 Читаем таблицы
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    text += "\n" + " | ".join(row_text)

    # PDF support removed
    # elif ext == 'pdf':
    #     with fitz.open(stream=file_bytes, filetype='pdf') as doc:
    #         for page in doc:
    #             text += page.get_text()

    else:
        raise ValueError("❌ Неподдерживаемый формат файла")

    return text

def extract_fields(text):
    def find(label):
        pattern = rf"{re.escape(label)}\s*(.*)"
        match = re.search(pattern, text)
        if match:
            return re.sub(r"^\\(.*?\\):?\\s*", "", match.group(1)).strip()
        return ""
    return {
        "initiator": find("Инициатор заявки (ФИО, должность):"),
        "description": find("Описание объекта закупки (кратко):"),
        "cost": find("Ориентировочная стоимость:"),
        "phone": find("Номер телефона инициатора заявки:"),
        "email": find("Адрес электронной почты:"),
        "department": find("Корпус, филиал:"),
    }
